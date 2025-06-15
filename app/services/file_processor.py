import os
import aiofiles
from typing import List, Dict, Any, BinaryIO
import PyPDF2
from docx import Document as DocxDocument
import tiktoken
from langchain_text_splitters import RecursiveCharacterTextSplitter
import structlog
from io import BytesIO
import hashlib

from app.core.config import settings
from app.models.rag import DocumentChunk

logger = structlog.get_logger()


class FileProcessor:
    def __init__(self):
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=self._token_length,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def _token_length(self, text: str) -> int:
        """Calculate token count for text"""
        return len(self.tokenizer.encode(text))
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        file_type: str
    ) -> List[DocumentChunk]:
        """Process uploaded file and return chunks"""
        
        # Extract text based on file type
        if file_type == 'pdf':
            text = await self._extract_pdf_text(file_content)
        elif file_type == 'docx':
            text = await self._extract_docx_text(file_content)
        elif file_type == 'txt':
            text = file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create DocumentChunk objects
        document_chunks = []
        for i, chunk_text in enumerate(chunks):
            chunk = DocumentChunk(
                text=chunk_text,
                metadata={
                    'filename': filename,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'tokens': self._token_length(chunk_text)
                }
            )
            document_chunks.append(chunk)
        
        await logger.ainfo(
            "file_processed",
            filename=filename,
            file_type=file_type,
            total_chunks=len(document_chunks),
            total_tokens=sum(c.metadata['tokens'] for c in document_chunks)
        )
        
        return document_chunks
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        text_parts = []
        
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    # Add page marker for better chunking
                    text_parts.append(f"\n[Page {page_num + 1}]\n{page_text}")
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error("pdf_extraction_error", error=str(e))
            raise ValueError(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        text_parts = []
        
        try:
            docx_file = BytesIO(content)
            doc = DocxDocument(docx_file)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error("docx_extraction_error", error=str(e))
            raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Fix common encoding issues
        replacements = {
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
            '–': '-',
            '—': '-',
            '…': '...',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text.strip()
    
    @staticmethod
    def generate_document_id(user_id: str, filename: str, content: bytes) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.sha256(content).hexdigest()[:8]
        return f"{user_id[:8]}_{content_hash}_{hashlib.sha256(filename.encode()).hexdigest()[:8]}"
    
    @staticmethod
    def validate_file(filename: str, file_size: int) -> tuple[bool, str]:
        """Validate file before processing"""
        # Check file extension
        file_ext = filename.lower().split('.')[-1]
        if file_ext not in settings.allowed_file_types:
            return False, f"File type '{file_ext}' not allowed. Allowed types: {', '.join(settings.allowed_file_types)}"
        
        # Check file size
        if file_size > settings.max_upload_size_bytes:
            max_mb = settings.max_upload_size_mb
            return False, f"File size exceeds {max_mb}MB limit"
        
        return True, file_ext